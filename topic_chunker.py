"""
topic_chunker.py (topic_chunks-only JSON output, embeddings always on, multi-file inputs)

What it does:
- Parses PDF/DOCX/TXT into structure-aware blocks
- Redacts/drops customer + PII + consultant/contact lines at parse time
- Labels each block (microsegment) with multiple topics using:
    heading prior + keyword hits + exclusions + embedding similarity (ALWAYS ON)
- Builds per-topic chunks by merging adjacent labeled segments
- Outputs JSON containing ONLY topic chunks (no microsegments, no blocks)

Usage:
  python topic_chunker.py --inputs a.pdf b.docx c.txt --output out.json --format auto

Dependencies:
  pip install pdfplumber python-docx sentence-transformers
"""

from __future__ import annotations

import argparse
import json
import math
import os
import re
from dataclasses import dataclass, asdict
from typing import Dict, List, Optional, Tuple, Any


# =========================
# Embeddings (always on)
# =========================
class Embedder:
    def embed(self, texts: List[str]) -> List[List[float]]:
        raise NotImplementedError


class SentenceTransformerEmbedder(Embedder):
    def __init__(self, model_name: str = "sentence-transformers/all-MiniLM-L6-v2"):
        try:
            from sentence_transformers import SentenceTransformer  # type: ignore
        except Exception as e:
            raise RuntimeError(
                "sentence-transformers not installed. Run: pip install sentence-transformers"
            ) from e
        self.model = SentenceTransformer(model_name)

    def embed(self, texts: List[str]) -> List[List[float]]:
        vecs = self.model.encode(texts, normalize_embeddings=True)
        return [v.tolist() for v in vecs]


def cosine(a: List[float], b: List[float]) -> float:
    dot = 0.0
    na = 0.0
    nb = 0.0
    for x, y in zip(a, b):
        dot += x * y
        na += x * x
        nb += y * y
    if na <= 0 or nb <= 0:
        return 0.0
    return dot / math.sqrt(na * nb)


# =========================
# Data types
# =========================
@dataclass
class Block:
    doc_id: str
    block_id: int
    type: str  # heading | paragraph | bullet | table_row
    text: str
    heading_path: List[str]
    page: Optional[int] = None
    section_number: Optional[str] = None


@dataclass
class LabeledSegment:
    segment_id: int
    block: Block
    labels: List[str]
    label_scores: Dict[str, float]
    evidence: Dict[str, List[str]]


@dataclass
class TopicChunk:
    topic: str
    chunk_text: str
    start_segment_id: int
    end_segment_id: int
    heading_path_start: List[str]
    heading_path_end: List[str]
    confidence: float
    evidence: List[str]
    input: str


# =========================
# Topics
# =========================
TOPICS = [
    "Use cases",
    "Architecture",
    "IGP, MPLS, SR",
    "BGP",
    "Services",
    "Hardware- Software",
    "QoS",
    "Scaling and Performance",
    "EMS-Paragon",
    "Miscellaneous",
]

TOPIC_CARDS: Dict[str, Dict[str, Any]] = {
    "Use cases": {
        "heading_keywords": ["use cases", "use-case", "scenario", "deployment", "migration strategy"],
        "strong": [
            "use case", "use cases", "deployment scenario", "customer scenario", "migration strategy",
            "greenfield", "brownfield", "rollout plan"
        ],
        "weak": ["scenario", "migration", "deployment", "workflow"],
        "exclude": ["bgp", "isis", "ospf", "srgb", "ti-lfa", "dscp", "policing", "shaping"],
    },
    "Architecture": {
        "heading_keywords": ["architecture", "solution overview", "design overview", "high level design"],
        "strong": [
            "architecture", "solution overview", "high level design", "control plane", "data plane",
            "topology", "network divided", "layers", "rings", "core", "aggregation"
        ],
        "weak": ["design", "overview", "component", "module", "interface"],
        "exclude": ["route reflector", "afi/safi", "isis instance", "ti-lfa", "dscp", "wred"],
    },
    "IGP, MPLS, SR": {
        "heading_keywords": ["igp", "mpls", "segment routing", "sr", "isis", "ospf", "ldp", "rsvp-te"],
        "strong": [
            "igp design", "isis", "is-is", "ospf", "ldp", "rsvp-te", "segment routing", "srgb",
            "sid", "segment list", "sr-te", "ti-lfa", "rlfa", "label stack", "mpls"
        ],
        "weak": ["prefix-sid", "adjacency sid", "node sid", "sr", "te", "bfd"],
        "exclude": ["afi/safi", "route reflector", "bgp-lu", "community", "dscp", "paragon"],
    },
    "BGP": {
        "heading_keywords": ["bgp", "evpn", "route reflector", "rr", "afi", "safi", "bgp-lu"],
        "strong": [
            "bgp design", "ibgp", "ebgp", "route reflector", "rr", "afi/safi", "address family",
            "path attribute", "community", "bgp-lu", "next-hop-self", "aigp", "evpn"
        ],
        "weak": ["policy", "neighbor", "session", "keepalive", "holdtime"],
        "exclude": ["isis", "ospf", "srgb", "ti-lfa", "dscp", "wred"],
    },
    "Services": {
        "heading_keywords": ["services", "l3vpn", "l2vpn", "evpn services", "vrf", "vlan", "enterprise"],
        "strong": [
            "services", "l3vpn", "l2vpn", "l2pw", "ill", "vrf", "vlan", "6vpe", "service chaining",
            "enterprise services", "subscriber services", "management vpn"
        ],
        "weak": ["vpn", "service", "tenant", "access"],
        "exclude": ["dscp", "wred", "queue", "policing", "shaping"],
    },
    "Hardware- Software": {
        "heading_keywords": ["hardware", "software", "platform", "line card", "asic", "firmware", "nos"],
        "strong": [
            "hardware", "software", "platform", "line card", "asic", "firmware", "nos",
            "minimum software version", "upgrade", "feature parity", "cpu", "memory",
            "device inventory", "inventory", "bom", "bill of materials", "juniper device"
        ],
        "weak": ["version", "release", "image", "upgrade path", "mx", "ptx", "acx", "qfx"],
        "exclude": [],
    },
    "QoS": {
        "heading_keywords": ["qos", "quality of service", "feature - qos"],
        "strong": [
            "qos", "quality of service", "dscp", "exp", "phb", "queue", "scheduler", "wred",
            "policing", "shaping", "marking", "classification"
        ],
        "weak": ["latency class", "traffic class", "buffer"],
        "exclude": [],
    },
    "Scaling and Performance": {
        "heading_keywords": ["scaling", "performance", "scale", "capacity", "convergence", "benchmark"],
        "strong": [
            "scaling", "performance", "scale", "scale needed", "capacity", "limits",
            "convergence", "failover", "throughput", "latency", "benchmark"
        ],
        "weak": ["cpu", "memory", "table", "churn"],
        "exclude": [],
    },
    "EMS-Paragon": {
        "heading_keywords": ["paragon", "ems", "nms", "controller", "pid", "ndd"],
        "strong": [
            "paragon", "ems", "nms", "controller", "pid", "ndd", "onboard device",
            "northbound", "telemetry integration", "workflow"
        ],
        "weak": ["api", "automation", "topology", "intent"],
        "exclude": [],
    },
    "Miscellaneous": {
        "heading_keywords": ["misc", "miscellaneous", "other", "open points"],
        "strong": ["miscellaneous", "mischalleneous", "open points", "parking lot"],
        "weak": ["note", "discussion", "tbd"],
        "exclude": [],
    },
}


# =========================
# Filtering (parse-time)
# =========================
DROP_LINE_PATTERNS = [
    r"^\s*Juniper\s+(Business\s+)?Use\s+Only\s*$",
    r"^\s*Juniper\s+Use\s+Only\s*$",
    r"^\s*Juniper\s+Business\s+Use\s+Only\s*$",
    r"^\s*Customer\s+ask\s*:\s*.*$",
    r"^\s*Customer\s+asks\s*:\s*.*$",
    r"^\s*Customer\s+requirement\s*:\s*.*$",
    r"^\s*(?:email|e-?mail|phone|mobile|tel|whatsapp)\s*[:\-]\s*.*$",
]
DROP_LINE_RE = re.compile("|".join(f"(?:{p})" for p in DROP_LINE_PATTERNS), re.IGNORECASE)

CUSTOMER_REDACT_PATTERNS = [
    r"\bVIL\b",
    r"\bVodafone\b",
    r"\bIdea\b",
    r"\bVodafone\s+Idea\b",
    r"\bVodafone\s+Idea\s+Limited\b",
]
CUSTOMER_REDACT_RE = re.compile("|".join(f"(?:{p})" for p in CUSTOMER_REDACT_PATTERNS), re.IGNORECASE)

EMAIL_RE = re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.IGNORECASE)
PHONE_RE = re.compile(r"(?<!\w)(?:\+?\d{1,3}[\s\-]?)?(?:\(?\d{2,5}\)?[\s\-]?)?\d{6,10}(?!\w)")
CONSULTANT_LINE_RE = re.compile(
    r"^\s*(consultant|consultants|author|prepared\s*by|created\s*by|presented\s*by|"
    r"owner|point\s*of\s*contact|poc|contact\s*person|from|to|cc)\s*[:\-]\s*.+$",
    re.IGNORECASE,
)

def redact_pii(text: str) -> str:
    text = EMAIL_RE.sub("[EMAIL]", text)
    text = PHONE_RE.sub("[PHONE]", text)
    return text

def redact_customer_and_people(text: str) -> str:
    text = CUSTOMER_REDACT_RE.sub("[CUSTOMER]", text)
    text = redact_pii(text)
    return text

def should_drop_line(text: str) -> bool:
    t = text.strip()
    if not t:
        return True
    if DROP_LINE_RE.match(t):
        return True
    if CONSULTANT_LINE_RE.match(t):
        return True
    if t in {"[EMAIL]", "[PHONE]", "[CUSTOMER]"}:
        return True
    if ("[EMAIL]" in t or "[PHONE]" in t) and re.fullmatch(r"[\[\]A-Z0-9\s:,\-+()./]*", t):
        return True
    return False


# =========================
# Structure heuristics
# =========================
INLINE_LOCAL_HEADINGS = [
    "IGP Design",
    "Segment Routing",
    "BGP Design",
    "Services",
    "Feature - QoS",
    "Use Cases",
    "Paragon EMS",
    "Hardware",
    "Software",
    "Scaling",
    "Performance",
]
INLINE_LOCAL_HEADING_RE = re.compile(
    r"^\s*(?:" + "|".join(re.escape(x) for x in INLINE_LOCAL_HEADINGS) + r")\s*:?\s*$",
    re.IGNORECASE,
)

NUMBERED_HEADING_RE = re.compile(r"^\s*(\d+(?:\.\d+)*)(?:\.)?\s+(.+?)\s*$")
BULLET_RE = re.compile(r"^\s*([•\-\*]|\d+\)|\d+\.)\s+")
WHITESPACE_RE = re.compile(r"\s+")
VERBISH = re.compile(r"\b(has|have|will|are|is|was|were|should|shall|can|may|must)\b", re.IGNORECASE)

def normalize_text(s: str) -> str:
    s = s.replace("\u00a0", " ")
    s = WHITESPACE_RE.sub(" ", s).strip()
    return s

def looks_like_numbered_heading(full_line: str, sec_no: str, title: str) -> bool:
    t = title.strip()
    if len(t.split()) > 10:
        return False
    if VERBISH.search(t):
        return False
    if "=" in t:
        return False
    if re.search(r"\b\d+\s*-\s*\d+\b", t):
        return False
    alpha = sum(ch.isalpha() for ch in t)
    if alpha >= max(4, int(0.5 * len(t))):
        return True
    return False


# =========================
# Extractors
# =========================
def extract_docx_blocks(path: str, doc_id: str) -> List[Block]:
    from docx import Document  # type: ignore

    doc = Document(path)
    blocks: List[Block] = []
    heading_path: List[str] = []
    block_id = 0

    def push_block(t: str, text: str, sec_no: Optional[str] = None):
        nonlocal block_id
        text = normalize_text(redact_customer_and_people(text))
        if not text or should_drop_line(text):
            return
        blocks.append(
            Block(
                doc_id=doc_id,
                block_id=block_id,
                type=t,
                text=text,
                heading_path=heading_path.copy(),
                section_number=sec_no,
            )
        )
        block_id += 1

    for p in doc.paragraphs:
        style = (p.style.name or "").lower() if p.style else ""
        raw = normalize_text(p.text)
        if not raw:
            continue

        raw = normalize_text(redact_customer_and_people(raw))
        if not raw or should_drop_line(raw):
            continue

        if "heading" in style:
            m = NUMBERED_HEADING_RE.match(raw)
            sec_no = m.group(1) if m else None
            level = 1
            m2 = re.search(r"heading\s+(\d+)", style)
            if m2:
                level = int(m2.group(1))
            while len(heading_path) >= level:
                heading_path.pop()
            heading_path.append(raw)
            push_block("heading", raw, sec_no=sec_no)
            continue

        if BULLET_RE.match(raw):
            push_block("bullet", raw)
        else:
            if INLINE_LOCAL_HEADING_RE.match(raw):
                push_block("heading", raw)
            else:
                push_block("paragraph", raw)

    for table in doc.tables:
        for row in table.rows:
            cells = [normalize_text(redact_customer_and_people(c.text)) for c in row.cells]
            cells = [c for c in cells if c and not should_drop_line(c)]
            row_text = " | ".join(cells)
            if row_text:
                push_block("table_row", row_text)

    return blocks

def extract_txt_blocks(path: str, doc_id: str) -> List[Block]:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        raw = f.read()

    lines = [line.rstrip() for line in raw.splitlines()]
    blocks: List[Block] = []
    heading_path: List[str] = []
    block_id = 0
    buf: List[str] = []

    def flush_paragraph():
        nonlocal block_id
        if not buf:
            return
        text = normalize_text(redact_customer_and_people(" ".join(buf)))
        buf.clear()
        if not text or should_drop_line(text):
            return
        blocks.append(
            Block(
                doc_id=doc_id,
                block_id=block_id,
                type="paragraph",
                text=text,
                heading_path=heading_path.copy(),
            )
        )
        block_id += 1

    def push_heading(text: str, sec_no: Optional[str]):
        nonlocal block_id
        text = normalize_text(redact_customer_and_people(text))
        if not text or should_drop_line(text):
            return
        blocks.append(
            Block(
                doc_id=doc_id,
                block_id=block_id,
                type="heading",
                text=text,
                heading_path=heading_path.copy(),
                section_number=sec_no,
            )
        )
        block_id += 1

    for line in lines:
        if not line.strip():
            flush_paragraph()
            continue

        text = normalize_text(redact_customer_and_people(line))
        if not text or should_drop_line(text):
            continue

        m = NUMBERED_HEADING_RE.match(text)
        if m:
            sec_no = m.group(1)
            title = m.group(2)
            if looks_like_numbered_heading(text, sec_no, title):
                flush_paragraph()
                depth = sec_no.count(".") + 1
                while len(heading_path) >= depth:
                    heading_path.pop()
                heading_path.append(text)
                push_heading(text, sec_no=sec_no)
                continue

        if INLINE_LOCAL_HEADING_RE.match(text):
            flush_paragraph()
            push_heading(text, sec_no=None)
            continue

        if BULLET_RE.match(text):
            flush_paragraph()
            blocks.append(
                Block(
                    doc_id=doc_id,
                    block_id=block_id,
                    type="bullet",
                    text=text,
                    heading_path=heading_path.copy(),
                )
            )
            block_id += 1
            continue

        buf.append(text)

    flush_paragraph()
    return blocks

def extract_pdf_blocks(path: str, doc_id: str) -> List[Block]:
    import pdfplumber  # type: ignore

    blocks: List[Block] = []
    heading_path: List[str] = []
    block_id = 0

    def push_block(t: str, text: str, page: int, sec_no: Optional[str] = None):
        nonlocal block_id
        text = normalize_text(redact_customer_and_people(text))
        if not text or should_drop_line(text):
            return
        blocks.append(
            Block(
                doc_id=doc_id,
                block_id=block_id,
                type=t,
                text=text,
                heading_path=heading_path.copy(),
                page=page,
                section_number=sec_no,
            )
        )
        block_id += 1

    with pdfplumber.open(path) as pdf:
        for pageno, page in enumerate(pdf.pages, start=1):
            words = page.extract_words(extra_attrs=["size", "fontname"])
            if not words:
                continue

            sizes = [round(float(w.get("size", 0.0))) for w in words if w.get("size")]
            body_size = 10
            if sizes:
                hist: Dict[int, int] = {}
                for s in sizes:
                    hist[s] = hist.get(s, 0) + 1
                body_size = max(hist.items(), key=lambda kv: kv[1])[0]

            words_sorted = sorted(words, key=lambda w: (round(w["top"], 1), w["x0"]))
            lines: List[List[dict]] = []
            current: List[dict] = []
            current_top: Optional[float] = None
            line_tol = 2.0

            for w in words_sorted:
                top = float(w["top"])
                if current_top is None or abs(top - current_top) <= line_tol:
                    current.append(w)
                    current_top = top if current_top is None else current_top
                else:
                    lines.append(current)
                    current = [w]
                    current_top = top
            if current:
                lines.append(current)

            line_objs: List[Tuple[str, float]] = []
            for line in lines:
                line = sorted(line, key=lambda w: w["x0"])
                text = " ".join([w["text"] for w in line])
                sizes_line = [float(w.get("size", body_size)) for w in line]
                rep_size = sum(sizes_line) / max(1, len(sizes_line))

                text = normalize_text(text)
                text = normalize_text(redact_customer_and_people(text))
                if not text or should_drop_line(text):
                    continue
                line_objs.append((text, rep_size))

            try:
                tables = page.extract_tables()
            except Exception:
                tables = []

            para_buf: List[str] = []

            def flush_para():
                nonlocal para_buf
                if not para_buf:
                    return
                txt = normalize_text(redact_customer_and_people(" ".join(para_buf)))
                para_buf = []
                if not txt or should_drop_line(txt):
                    return
                if BULLET_RE.match(txt):
                    push_block("bullet", txt, pageno)
                else:
                    push_block("paragraph", txt, pageno)

            for text, size in line_objs:
                if not text:
                    flush_para()
                    continue

                m = NUMBERED_HEADING_RE.match(text)
                if m:
                    sec_no = m.group(1)
                    title = m.group(2)
                    if looks_like_numbered_heading(text, sec_no, title):
                        flush_para()
                        depth = sec_no.count(".") + 1
                        while len(heading_path) >= depth:
                            heading_path.pop()
                        heading_path.append(text)
                        push_block("heading", text, pageno, sec_no=sec_no)
                        continue

                if INLINE_LOCAL_HEADING_RE.match(text):
                    flush_para()
                    push_block("heading", text, pageno, sec_no=None)
                    continue

                if size >= body_size + 2 and len(text.split()) <= 12:
                    flush_para()
                    heading_path.append(text)
                    if len(heading_path) > 6:
                        heading_path[:] = heading_path[-6:]
                    push_block("heading", text, pageno, sec_no=None)
                    continue

                if BULLET_RE.match(text):
                    flush_para()
                    push_block("bullet", text, pageno)
                    continue

                para_buf.append(text)

            flush_para()

            for t in tables or []:
                for row in t:
                    if not row:
                        continue
                    row_cells: List[str] = []
                    for c in row:
                        if c is None:
                            continue
                        cell = normalize_text(redact_customer_and_people(str(c)))
                        if cell and not should_drop_line(cell):
                            row_cells.append(cell)
                    row_text = " | ".join(row_cells)
                    if row_text:
                        push_block("table_row", row_text, pageno)

    return blocks

def extract_blocks(path: str, fmt: str = "auto") -> List[Block]:
    doc_id = os.path.basename(path)
    ext = os.path.splitext(path)[1].lower()
    if fmt == "auto":
        if ext == ".pdf":
            fmt = "pdf"
        elif ext == ".docx":
            fmt = "docx"
        elif ext in [".txt", ".md"]:
            fmt = "txt"
        else:
            raise ValueError(f"Unknown extension for auto mode: {ext}")

    if fmt == "pdf":
        return extract_pdf_blocks(path, doc_id)
    if fmt == "docx":
        return extract_docx_blocks(path, doc_id)
    if fmt == "txt":
        return extract_txt_blocks(path, doc_id)
    raise ValueError(f"Unsupported format: {fmt}")


# =========================
# Labeling
# =========================
def _find_evidence(text: str, patterns: List[str]) -> List[str]:
    lowered = text.lower()
    found: List[str] = []

    for p in patterns:
        pl = p.lower().strip()
        if not pl:
            continue

        # phrase: substring ok
        if " " in pl:
            if pl in lowered:
                found.append(p)
            continue

        # single token: strict boundaries (prevents "ill" in "will")
        token_re = rf"(?<!\w){re.escape(pl)}(?!\w)"
        if re.search(token_re, lowered):
            found.append(p)

    uniq: List[str] = []
    seen = set()
    for x in found:
        k = x.lower()
        if k not in seen:
            uniq.append(x)
            seen.add(k)
    return uniq

def label_segments(
    blocks: List[Block],
    embedder: Embedder,
    max_labels_per_segment: int = 3,
    emb_sim_threshold: float = 0.55,
    emb_sim_high: float = 0.68,
) -> List[LabeledSegment]:
    topic_repr_texts: Dict[str, str] = {}
    for t in TOPICS:
        card = TOPIC_CARDS[t]
        topic_repr_texts[t] = (
            f"Topic: {t}\n"
            f"Description keywords: {', '.join(card.get('strong', [])[:12])}\n"
            f"Also: {', '.join(card.get('weak', [])[:12])}\n"
        )

    topic_vecs_list = embedder.embed([topic_repr_texts[t] for t in TOPICS])
    topic_vecs: Dict[str, List[float]] = {t: v for t, v in zip(TOPICS, topic_vecs_list)}

    seg_inputs = []
    for b in blocks:
        hp = " > ".join(b.heading_path[-3:]) if b.heading_path else ""
        seg_inputs.append(f"{hp}\n{b.text}")
    seg_vecs = embedder.embed(seg_inputs)

    labeled: List[LabeledSegment] = []

    for i, b in enumerate(blocks):
        text = b.text
        hp_text = " > ".join(b.heading_path).lower()

        scores: Dict[str, float] = {t: 0.0 for t in TOPICS}
        evidence: Dict[str, List[str]] = {t: [] for t in TOPICS}
        sims: Dict[str, float] = {}

        for t in TOPICS:
            card = TOPIC_CARDS[t]

            heading_hits = _find_evidence(hp_text, card.get("heading_keywords", []))
            if heading_hits:
                scores[t] += 2.5
                evidence[t].extend(heading_hits)

            strong_hits = _find_evidence(text, card.get("strong", []))
            weak_hits = _find_evidence(text, card.get("weak", []))
            if strong_hits:
                scores[t] += 1.2 * min(3, len(strong_hits))
                evidence[t].extend(strong_hits)
            if weak_hits:
                scores[t] += 0.4 * min(4, len(weak_hits))

            excl_hits = _find_evidence(text, card.get("exclude", []))
            if excl_hits:
                scores[t] -= 1.5 * min(2, len(excl_hits))

            sim = cosine(seg_vecs[i], topic_vecs[t])
            sims[t] = sim
            scores[t] += 1.0 * sim

        chosen: List[Tuple[str, float]] = []
        for t in TOPICS:
            if t == "Miscellaneous":
                continue

            card = TOPIC_CARDS[t]
            hp_hits = _find_evidence(hp_text, card.get("heading_keywords", []))
            strong_hits = _find_evidence(text, card.get("strong", []))
            weak_hits = _find_evidence(text, card.get("weak", []))

            sim = sims.get(t, 0.0)
            heading_match = len(hp_hits) > 0

            add = False
            if heading_match:
                add = True
            elif len(strong_hits) >= 2:
                add = True
            elif len(strong_hits) >= 1 and sim >= emb_sim_threshold:
                add = True
            elif sim >= emb_sim_high and len(weak_hits) >= 1:
                add = True

            if add:
                if t in ["Architecture", "Use cases", "Services"] and not heading_match and len(strong_hits) == 0:
                    add = False

            if add and scores[t] > 0.2:
                chosen.append((t, scores[t]))
                if weak_hits:
                    evidence[t].extend(weak_hits)

        chosen.sort(key=lambda x: x[1], reverse=True)
        labels = [t for t, _ in chosen[:max_labels_per_segment]]

        if not labels:
            labels = ["Miscellaneous"]
            evidence["Miscellaneous"] = _find_evidence(text, TOPIC_CARDS["Miscellaneous"]["strong"])

        labeled.append(
            LabeledSegment(
                segment_id=i,
                block=b,
                labels=labels,
                label_scores={t: scores[t] for t in labels},
                evidence={t: evidence[t] for t in labels},
            )
        )

    return labeled

def context_inherit_labels(
    labeled: List[LabeledSegment],
    short_token_threshold: int = 25,
    inherit_min_score: float = 2.5,
    max_inherit_topics: int = 2,
) -> List[LabeledSegment]:
    def tok(s: str) -> int:
        return max(1, len(s.split()))

    heading_dominant: Dict[str, Tuple[str, float]] = {}
    for seg in labeled:
        if not seg.block.heading_path:
            continue
        hp_key = " > ".join(seg.block.heading_path)
        best: Optional[Tuple[str, float]] = None
        for t, sc in seg.label_scores.items():
            if t == "Miscellaneous":
                continue
            if best is None or sc > best[1]:
                best = (t, sc)
        if best and best[1] >= inherit_min_score:
            if hp_key not in heading_dominant or best[1] > heading_dominant[hp_key][1]:
                heading_dominant[hp_key] = best

    for i, seg in enumerate(labeled):
        if seg.labels != ["Miscellaneous"]:
            continue
        if tok(seg.block.text) >= short_token_threshold:
            continue

        candidates: List[Tuple[str, float]] = []

        if i > 0:
            prev = labeled[i - 1]
            for t, sc in prev.label_scores.items():
                if t != "Miscellaneous" and sc >= inherit_min_score:
                    candidates.append((t, sc))

        if seg.block.heading_path:
            hp_key = " > ".join(seg.block.heading_path)
            if hp_key in heading_dominant:
                candidates.append(heading_dominant[hp_key])

        if not candidates:
            continue

        candidates.sort(key=lambda x: x[1], reverse=True)
        new_topics: List[str] = []
        seen = set()
        for t, _ in candidates:
            if t not in seen:
                new_topics.append(t)
                seen.add(t)
            if len(new_topics) >= max_inherit_topics:
                break

        seg.labels = new_topics
        seg.label_scores = {t: max(seg.label_scores.get(t, 0.0), inherit_min_score) for t in new_topics}
        seg.evidence = {t: ["(inherited from context)"] for t in new_topics}

    return labeled


# =========================
# Chunking
# =========================
def estimate_tokens(text: str) -> int:
    return max(1, len(text.split()))

def build_topic_chunks(
    labeled_segments: List[LabeledSegment],
    input_path: str,
    max_chunk_tokens: int = 900,
    break_on_new_numbered_section: bool = True,
) -> List[TopicChunk]:
    segs = labeled_segments
    topic_to_ids: Dict[str, List[int]] = {t: [] for t in TOPICS}
    for s in segs:
        for t in s.labels:
            topic_to_ids[t].append(s.segment_id)

    chunks: List[TopicChunk] = []

    for topic in TOPICS:
        ids = topic_to_ids.get(topic, [])
        if not ids:
            continue

        i = 0
        while i < len(ids):
            start_id = ids[i]
            end_id = start_id

            parts = [segs[start_id].block.text]
            evid: List[str] = []
            confs = [segs[start_id].label_scores.get(topic, 0.0)]
            evid.extend(segs[start_id].evidence.get(topic, []))

            current_tokens = estimate_tokens(parts[0])

            def section_no(seg_id: int) -> Optional[str]:
                return segs[seg_id].block.section_number

            j = i + 1
            while j < len(ids):
                nxt = ids[j]
                if nxt != end_id + 1:
                    break

                if break_on_new_numbered_section:
                    prev_sec = section_no(end_id)
                    nxt_sec = section_no(nxt)
                    if nxt_sec and prev_sec and nxt_sec != prev_sec:
                        break

                nxt_text = segs[nxt].block.text
                nxt_tokens = estimate_tokens(nxt_text)
                if current_tokens + nxt_tokens > max_chunk_tokens:
                    break

                parts.append(nxt_text)
                current_tokens += nxt_tokens
                end_id = nxt
                confs.append(segs[nxt].label_scores.get(topic, 0.0))
                evid.extend(segs[nxt].evidence.get(topic, []))
                j += 1

            chunk_text = "\n".join(parts).strip()
            confidence = float(min(confs)) if confs else 0.0

            ev_uniq: List[str] = []
            seen = set()
            for e in evid:
                k = e.lower()
                if k not in seen:
                    ev_uniq.append(e)
                    seen.add(k)

            chunks.append(
                TopicChunk(
                    topic=topic,
                    chunk_text=chunk_text,
                    start_segment_id=start_id,
                    end_segment_id=end_id,
                    heading_path_start=segs[start_id].block.heading_path,
                    heading_path_end=segs[end_id].block.heading_path,
                    confidence=confidence,
                    evidence=ev_uniq,
                    input=input_path,
                )
            )

            i = j

    # keep output readable in original doc order
    chunks.sort(key=lambda c: (c.input, c.start_segment_id, TOPICS.index(c.topic)))
    return chunks


# =========================
# Orchestration (topic_chunks only)
# =========================
def process_file(
    input_path: str,
    embedder: Embedder,
    fmt: str = "auto",
    max_labels_per_segment: int = 3,
    max_chunk_tokens: int = 900,
    enable_context_inheritance: bool = True,
) -> List[Dict[str, Any]]:
    blocks = extract_blocks(input_path, fmt=fmt)
    labeled = label_segments(
        blocks,
        embedder=embedder,
        max_labels_per_segment=max_labels_per_segment,
    )

    if enable_context_inheritance:
        labeled = context_inherit_labels(labeled)

    topic_chunks = build_topic_chunks(
        labeled,
        input_path=input_path,
        max_chunk_tokens=max_chunk_tokens,
        break_on_new_numbered_section=True,
    )

    # JSON output: topic chunks only
    return [asdict(c) for c in topic_chunks]


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument(
        "--inputs",
        nargs="+",
        required=True,
        help="One or more input files (.pdf/.docx/.txt). Example: --inputs a.pdf b.docx c.txt",
    )
    ap.add_argument("--output", required=True, help="Output JSON path")
    ap.add_argument("--format", default="auto", choices=["auto", "pdf", "docx", "txt"])
    ap.add_argument("--embedding-model", default="sentence-transformers/all-MiniLM-L6-v2")
    ap.add_argument("--max-labels", type=int, default=3)
    ap.add_argument("--max-chunk-tokens", type=int, default=900)
    ap.add_argument("--no-context-inheritance", action="store_true")
    args = ap.parse_args()

    embedder = SentenceTransformerEmbedder(args.embedding_model)

    all_chunks: List[Dict[str, Any]] = []
    for inp in args.inputs:
        chunks = process_file(
            input_path=inp,
            embedder=embedder,
            fmt=args.format,
            max_labels_per_segment=args.max_labels,
            max_chunk_tokens=args.max_chunk_tokens,
            enable_context_inheritance=(not args.no_context_inheritance),
        )
        all_chunks.extend(chunks)

    # Output only the chunks list
    from pathlib import Path

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(all_chunks, f, ensure_ascii=False, indent=2)


    print(f"Wrote {args.output}")
    print(f"Total topic chunks: {len(all_chunks)}")


if __name__ == "__main__":
    main()
